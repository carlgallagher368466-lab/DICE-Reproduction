# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
import re
import shutil
import subprocess

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "five_star_final_report_source.md"
DOCX = ROOT / "DICE_解耦学习推荐系统去偏_研究复现报告_五星完整版.docx"
PDF = ROOT / "DICE_解耦学习推荐系统去偏_研究复现报告_五星完整版.pdf"
RENDER_DIR = ROOT / "render_pages"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def paragraph(doc, text="", *, size=11, bold=False, color=None, align=None, before=0, after=6, line=1.10, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        write_inline(p, text, size=size, bold=bold, color=color, italic=italic)
    return p


def write_inline(paragraph_obj, text, *, size=11, bold=False, color=None, italic=False):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        part_bold = bold
        part_color = color
        content = part
        if part.startswith("**") and part.endswith("**"):
            part_bold = True
            content = part[2:-2]
        elif part.startswith("`") and part.endswith("`"):
            content = part[1:-1]
            part_color = DARK
        r = paragraph_obj.add_run(content)
        set_run_font(r, size=size, bold=part_bold, color=part_color, italic=italic)


def heading(doc, text, level):
    style = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}.get(level, 6))
    p.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}.get(level, 4))
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK)
    return p


def cell_text(cell, text, *, bold=False, size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    write_inline(p, str(text), size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    small = len(headers) >= 7
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, size=7.0 if small else 8.4)
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            value = str(value)
            align = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 16 and not re.match(r"^-?\d+(\.\d+)?$", value) else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(cells[i], value, size=6.8 if small else 8.0, align=align)
    paragraph(doc, "", after=2)
    return table


def parse_md_table(lines, start):
    headers = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    rows = []
    i = start + 2
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return headers, rows, i


def add_image(doc, rel_path, caption):
    path = ROOT / rel_path
    if not path.exists():
        paragraph(doc, f"[图缺失：{rel_path}]", color=RGBColor(155, 28, 28))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.1))
    paragraph(doc, caption, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("智能商务期末作业 | DICE 研究复现报告")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("解耦学习与因果嵌入推荐系统去偏")
    set_run_font(r, size=9, color=MUTED)
    return doc


def build_docx():
    doc = setup_document()
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            if first_title:
                paragraph(doc, "研究复现报告", size=12, bold=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=18)
                paragraph(doc, "解耦学习与因果嵌入的推荐系统去偏方法", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
                paragraph(doc, "研究复现报告", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
                first_title = False
            else:
                heading(doc, line[2:].strip(), 1)
            i += 1
            continue
        if line.startswith("## "):
            heading(doc, line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            heading(doc, line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("|"):
            headers, rows, i = parse_md_table(lines, i)
            add_table(doc, headers, rows)
            continue
        if line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                add_image(doc, match.group(2), match.group(1))
            i += 1
            continue
        if line.startswith("- "):
            while i < len(lines) and lines[i].startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(4)
                write_inline(p, lines[i][2:].strip(), size=10.5)
                i += 1
            continue
        if re.match(r"^\d+\. ", line):
            while i < len(lines) and re.match(r"^\d+\. ", lines[i]):
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(4)
                write_inline(p, re.sub(r"^\d+\. ", "", lines[i]).strip(), size=10.5)
                i += 1
            continue
        if line.startswith("**副标题"):
            paragraph(doc, line.replace("**", ""), size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
        elif line.startswith("**课程") or line.startswith("**复现论文") or line.startswith("**代码基础") or line.startswith("**完成时间"):
            paragraph(doc, line.replace("**", ""), size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        elif line.startswith("学生姓名"):
            paragraph(doc, line, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=12)
            doc.add_page_break()
        else:
            paragraph(doc, line)
        i += 1
    doc.save(DOCX)
    return DOCX


def export_pdf():
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(ROOT), str(DOCX)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    generated = ROOT / (DOCX.stem + ".pdf")
    if generated.exists() and generated != PDF:
        generated.replace(PDF)
    return PDF if PDF.exists() else generated


if __name__ == "__main__":
    path = build_docx()
    print(path)
    pdf = export_pdf()
    if pdf:
        print(pdf)
    else:
        print("PDF export skipped: soffice/libreoffice not found")
