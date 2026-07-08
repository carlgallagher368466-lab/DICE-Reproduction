# -*- coding: utf-8 -*-
from pathlib import Path
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FormatStrFormatter

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path("D:/codex/智能商务/outputs/final_report")
SOURCE = ROOT / "final_report_source.md"
FIG = ROOT / "figures"
FIG.mkdir(parents=True, exist_ok=True)
DOCX = ROOT / "DICE_解耦学习推荐系统去偏_研究复现报告_四星完整扩展版.docx"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"


def make_figures():
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    def grouped(path, title, labels, series, ylabel="Recall@20"):
        fig, ax = plt.subplots(figsize=(7.2, 4.0), dpi=180)
        x = list(range(len(labels)))
        width = 0.24
        colors = ["#4C78A8", "#F58518", "#54A24B"]
        for i, (name, vals) in enumerate(series):
            ax.bar([v + (i - 1) * width for v in x], vals, width=width, label=name, color=colors[i])
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_ylabel(ylabel)
        ax.set_title(title, fontsize=12)
        ax.yaxis.set_major_formatter(FormatStrFormatter("%.2f"))
        ax.grid(axis="y", alpha=0.25)
        ax.legend(frameon=False, ncol=3, fontsize=9)
        fig.tight_layout()
        fig.savefig(path, bbox_inches="tight")
        plt.close(fig)

    grouped(
        FIG / "fig1_neumf_recall20.png",
        "NeuMF 系列在三类测试分布上的 Recall@20",
        ["Longtail", "Uniform", "Head"],
        [
            ("NeuMF", [0.1491, 0.2893, 0.3548]),
            ("NeuMFIPS", [0.1169, 0.1242, 0.1100]),
            ("NeuMFDICE", [0.1551, 0.1625, 0.1000]),
        ],
    )
    grouped(
        FIG / "fig2_ngcf_recall20.png",
        "NGCF 系列在三类测试分布上的 Recall@20",
        ["Longtail", "Uniform", "Head"],
        [
            ("NGCF", [0.1393, 0.3156, 0.4016]),
            ("NGCFIPS", [0.1544, 0.1901, 0.1721]),
            ("NGCFDICE", [0.1769, 0.2286, 0.1596]),
        ],
    )
    grouped(
        FIG / "fig5_vae_recall20.png",
        "VAE 系列在三类测试分布上的 Recall@20",
        ["Longtail", "Uniform", "Head"],
        [
            ("VAE", [0.1136, 0.2652, 0.3340]),
            ("VAEIPS", [0.0860, 0.1003, 0.0824]),
            ("VAEDICE", [0.1664, 0.1774, 0.1102]),
        ],
    )
    grouped(
        FIG / "fig3_ablation_recall20.png",
        "NeuMF-DICE 消融实验指标对比",
        ["baseline", "no-dis", "no-adaptive", "low-dis", "high-dis"],
        [
            ("Recall@20", [0.1551, 0.1583, 0.1465, 0.1569, 0.1427]),
            ("NDCG@20", [0.1038, 0.1066, 0.0979, 0.1054, 0.0935]),
            ("Hit@20", [0.5341, 0.5413, 0.5194, 0.5448, 0.5034]),
        ],
        ylabel="Metric",
    )

    labels = ["no-dis", "no-adaptive", "low-dis", "high-dis"]
    avg_pop = [1103.3716, 1027.4270, 982.3677, 1194.6630]
    coverage = [0.4571, 0.7595, 0.5206, 0.5489]
    fig, ax1 = plt.subplots(figsize=(7.2, 4.0), dpi=180)
    x = list(range(len(labels)))
    ax1.bar([i - 0.18 for i in x], avg_pop, width=0.36, label="AvgPop@20", color="#E45756")
    ax1.set_ylabel("AvgPop@20")
    ax1.set_xticks(x)
    ax1.set_xticklabels(labels)
    ax1.grid(axis="y", alpha=0.25)
    ax2 = ax1.twinx()
    ax2.bar([i + 0.18 for i in x], coverage, width=0.36, label="Coverage@20", color="#72B7B2")
    ax2.set_ylabel("Coverage@20")
    ax1.set_title("消融实验中的流行度依赖与覆盖率")
    lines, labs = ax1.get_legend_handles_labels()
    lines2, labs2 = ax2.get_legend_handles_labels()
    ax1.legend(lines + lines2, labs + labs2, frameon=False, loc="upper center", ncol=2, fontsize=9)
    fig.tight_layout()
    fig.savefig(FIG / "fig4_popularity_coverage.png", bbox_inches="tight")
    plt.close(fig)


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def para(doc, text="", size=11, bold=False, color=None, align=None, before=0, after=6, line=1.10, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK)
    return p


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, bold=False, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    small = len(headers) >= 7
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, size=7.5 if small else 8.8)
        shade(table.rows[0].cells[i], LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(cells[i], value, size=7.2 if small else 8.3, align=align)
    para(doc, "", after=2)
    return table


def parse_md_table(lines, start):
    headers = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    rows = []
    i = start + 2
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return headers, rows, i


def add_markdown_paragraph(doc, text):
    if text.startswith("**") and text.endswith("**") and "：" not in text[:8]:
        para(doc, text.strip("*"), bold=True)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.10
        # Simple bold segments.
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if not part:
                continue
            bold = part.startswith("**") and part.endswith("**")
            content = part[2:-2] if bold else part
            r = p.add_run(content)
            set_run_font(r, size=11, bold=bold)


def build_docx():
    make_figures()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = FONT_CN
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("智能商务期末作业 | DICE 研究复现报告")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("DICE 研究复现报告")
    set_run_font(r, size=9, color=MUTED)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            if first_title:
                para(doc, "研究复现报告", size=12, bold=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=18)
                para(doc, "解耦学习与因果嵌入的推荐系统去偏方法", size=19, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line=1.15)
                para(doc, "研究复现报告", size=19, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, line=1.15)
                first_title = False
            else:
                heading(doc, line[2:].strip(), 1)
            i += 1
            continue
        if line.startswith("## "):
            heading(doc, line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            heading(doc, line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("|"):
            headers, rows, i = parse_md_table(lines, i)
            add_table(doc, headers, rows)
            continue
        if line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                caption, rel = m.group(1), m.group(2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(ROOT / rel), width=Inches(6.1))
                para(doc, caption, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
            i += 1
            continue
        if line.startswith("- "):
            while i < len(lines) and lines[i].startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(lines[i][2:].strip())
                set_run_font(r, size=10.5)
                i += 1
            continue
        if re.match(r"^\d+\. ", line):
            while i < len(lines) and re.match(r"^\d+\. ", lines[i]):
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(re.sub(r"^\d+\. ", "", lines[i]).strip())
                set_run_font(r, size=10.5)
                i += 1
            continue
        if line.startswith("**副标题"):
            para(doc, line.replace("**", ""), size=13, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
        elif line.startswith("**课程") or line.startswith("**论文") or line.startswith("**代码") or line.startswith("**完成时间"):
            para(doc, line.replace("**", ""), size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        elif line.startswith("学生姓名"):
            para(doc, line, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=12)
            doc.add_page_break()
        else:
            add_markdown_paragraph(doc, line)
        i += 1

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build_docx()
